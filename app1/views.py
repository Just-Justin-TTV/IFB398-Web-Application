# app1/views.py 

import json
import logging
import re
from io import BytesIO
from datetime import timedelta
from django.template.loader import get_template
from decimal import Decimal, InvalidOperation
from typing import Optional, Any, List
from django.contrib import messages
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.contrib.auth import authenticate, login, logout, update_session_auth_hash
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User
from django.db import connection, transaction
from django.db.models import Q, Avg, Count
from django.db.models.functions import Coalesce, ExtractYear
from django.http import (
    HttpRequest,
    HttpResponse,
    HttpResponseBadRequest,
    JsonResponse,
)
from django.shortcuts import get_object_or_404, redirect, render
from django.urls import reverse
from django.utils import timezone
from django.utils.text import slugify
from django.views.decorators.http import require_GET, require_POST

from .models import (
    ClassTargets,
    InterventionDependencies,
    InterventionEffects,
    Interventions,
    Metrics,
    User as AppUser,
    InterventionSelection,  # Stores selected interventions per project
)

logger = logging.getLogger(__name__)

# =========================
# Helpers
# =========================

def _resolve_app_user(request: HttpRequest) -> Optional[AppUser]:
    """
    Map the Django auth user to the AppUser row by username/email.
    Returns None if user is not authenticated or no matching AppUser found.
    """
    user = getattr(request, "user", None)
    if not getattr(user, "is_authenticated", False):
        return None
    if getattr(user, "username", None):
        hit = AppUser.objects.filter(username=user.username).first()
        if hit:
            return hit
    if getattr(user, "email", None):
        hit = AppUser.objects.filter(email=user.email).first()
        if hit:
            return hit
    return None


def _num(value: Any, default: Optional[float] = None) -> Optional[float]:
    """
    Convert various string/number formats to float.
    Handles units like AUD, %, k for thousand, commas, and dash replacements.
    """
    if value is None:
        return default
    s = str(value).strip().lower()
    s = (
        s.replace("aud", "")
        .replace(",", "")
        .replace("k", "000")
        .replace("–", "-")
        .replace("%", "")
        .strip()
    )
    try:
        return float(s)
    except Exception:
        return default


def _to_int(value: Any) -> Optional[int]:
    """
    Safely convert input to int.
    Returns None for empty strings, None, or 'null'.
    """
    if value in (None, "", "null"):
        return None
    try:
        return int(str(value).strip())
    except (ValueError, TypeError):
        return None


NUM_RE = re.compile(r"[^0-9\.\-]")  # Regex to keep digits, dot, and minus only

def _to_dec(value: Any, *, default: Optional[Decimal] = None) -> Optional[Decimal]:
    """
    Safely convert user input to Decimal.
    Cleans input by removing commas, units, %, spaces.
    Returns default for invalid or special values like NaN, inf, or empty strings.
    """
    if value in (None, "", "null"):
        return default
    s = str(value).strip()
    s_lower = s.lower()
    if s_lower in ("nan", "+nan", "-nan", "inf", "+inf", "-inf"):
        return default
    cleaned = NUM_RE.sub("", s)
    if cleaned in ("", "-", ".", "-.", ".-"):
        return default
    try:
        return Decimal(cleaned)
    except (InvalidOperation, ValueError, TypeError):
        return default


def _unique_project_code(project_name: str) -> str:
    """
    Generate a unique slug for Metrics.project_code.
    If the slug exists, append a number suffix: project, project-2, project-3...
    """
    base = slugify(project_name) or "project"
    code = base
    n = 2
    while Metrics.objects.filter(project_code=code).exists():
        code = f"{base}-{n}"
        n += 1
    return code


def _get_current_metric(request) -> Metrics:
    """
    Determine which Metrics row should be used for calculations.
    Priority:
      1) metrics_id from POST/GET
      2) metrics_id stored in session
      3) latest project for this user
      4) fallback: create a new Metrics row if none exists
    """
    app_user = _resolve_app_user(request)
    metrics_id = (
        request.POST.get("metrics_id")
        or request.GET.get("metrics_id")
        or request.session.get("metrics_id")
    )

    m = Metrics.objects.filter(id=metrics_id).first() if metrics_id else None
    if not m and app_user:
        m = (
            Metrics.objects.filter(user=app_user)
            .order_by("-updated_at", "-created_at")
            .first()
        )
    if not m:
        m = Metrics.objects.create(user=app_user if app_user else None)

    # Store current metrics ID in session
    request.session["metrics_id"] = m.id
    request.session.modified = True
    return m


# =========================
# Create & Edit Project
# =========================

@login_required(login_url='login')
def create_project(request: HttpRequest):
    """
    Handle creating a new project (Metrics row).
    POST -> create Metrics with project details and save session info.
    GET -> render create_project.html form.
    """
    if request.method == "POST":
        project_name = (request.POST.get("project_name") or "").strip()
        project_location = (request.POST.get("location") or request.POST.get("project_location") or "").strip()
        project_type = (request.POST.get("project_type") or "").strip()

        if not project_name:
            messages.error(request, "Project name is required.")
            return render(request, "create_project.html")

        code = _unique_project_code(project_name)
        m = Metrics.objects.create(
            user=_resolve_app_user(request),
            project_code=code,
            project_name=project_name,
            location=project_location,
            building_type=project_type,
        )

        # Save project IDs in session
        ids = list(request.session.get("my_project_ids", []))
        if m.id not in ids:
            ids.append(m.id)
        request.session["my_project_ids"] = ids
        request.session["metrics_id"] = m.id
        request.session.modified = True

        return redirect("carbon")

    return render(request, "create_project.html")


@login_required(login_url='login')
def metrics_edit(request, pk: int):
    """
    Edit basic project info stored in Metrics.
    GET -> render prefilled form
    POST -> save changes and redirect to Projects
    """
    m = get_object_or_404(Metrics, pk=pk)

    if request.method == "POST":
        m.project_name = (request.POST.get("project_name") or m.project_name or "").strip()
        m.location = (request.POST.get("location") or m.location or "").strip()
        m.building_type = (request.POST.get("project_type") or m.building_type or "").strip()
        m.save()

        # Update session with current Metrics ID
        request.session["metrics_id"] = m.id
        request.session.modified = True

        return redirect("projects")

    return render(request, "create_project.html", {"m": m, "is_edit": True})


# =========================
# Interventions API
# =========================

CLASS_ALIASES = {
    "carbon": ["carbon", "carbon emissions", "operating carbon", "operational carbon", "embodied carbon"],
    "health": ["health", "health & wellbeing", "health and wellbeing"],
    "water": ["water", "water use", "water efficiency"],
    "circular": ["circular", "circular economy"],
    "resilience": ["resilience"],
    "biodiversity": ["biodiversity"],
    "value": ["value", "value & cost", "value and cost"],
}

@require_GET
def interventions_api(request):
    """
    Return interventions as JSON, optionally filtered by class/theme.
    Includes current project metrics (if metrics_id in session).
    """
    ui_key = (request.GET.get("cls") or "").strip().lower()

    # Get metrics for current project
    metrics = {"gifa_m2": 0, "building_footprint_m2": 0}
    metrics_id = request.session.get("metrics_id")
    if metrics_id:
        try:
            metric_obj = Metrics.objects.filter(id=metrics_id).first()
            if metric_obj:
                metrics["gifa_m2"] = float(metric_obj.gifa_m2 or 0)
                metrics["building_footprint_m2"] = float(metric_obj.building_footprint_m2 or 0)
        except Exception:
            logger.exception("Error fetching metrics for metrics_id=%s", metrics_id)

    # Fetch interventions with optional filtering by class
    try:
        with connection.cursor() as cur:
            desc = connection.introspection.get_table_description(cur, "Interventions")
            colnames = [getattr(c, "name", getattr(c, "column_name", "")) for c in desc]

        select_cols = ", ".join(f'"{c}"' if c.lower() == "class" else c for c in colnames)
        sql = f'SELECT {select_cols} FROM "Interventions"'
        params = []

        if ui_key:
            terms = CLASS_ALIASES.get(ui_key, [ui_key])
            like_parts = []
            for t in terms:
                like_parts.append('LOWER("class") LIKE LOWER(%s)')
                params.append(f"%{t}%")
            sql += " WHERE " + " OR ".join(like_parts)

        sql += " ORDER BY theme, name"

        items = []
        with connection.cursor() as cur:
            cur.execute(sql, params)
            cols = [c[0] for c in cur.description]
            for row in cur.fetchall():
                obj = dict(zip(cols, row))
                items.append(
                    {
                        "id": obj.get("id"),
                        "name": obj.get("name") or f"Intervention #{obj.get('id')}",
                        "theme": obj.get("theme") or "",
                        "description": obj.get("description") or "",
                        "cost_level": float(obj.get("cost_level") or 0),
                        "intervention_rating": float(obj.get("intervention_rating") or 0),
                        "gifa_m2": metrics.get("gifa_m2", 0),
                        "building_footprint_m2": metrics.get("building_footprint_m2", 0),
                    }
                )
    except Exception:
        logger.exception("Error fetching interventions from DB")
        items = []

    return JsonResponse({"items": items})


# =========================
# Save Metrics
# =========================

@require_POST
@login_required(login_url='login')
def save_metrics(request: HttpRequest) -> JsonResponse:
    """
    Save building metrics for the current project.
    Handles Decimal, Integer, Boolean, and string fields.
    Stores Total Budget and updates session with current Metrics ID.
    """
    try:
        payload = json.loads(request.body.decode("utf-8") or "{}")
    except json.JSONDecodeError:
        return HttpResponseBadRequest("Invalid JSON")

    # Find or create Metrics instance
    metrics_id = payload.get("metrics_id") or request.session.get("metrics_id")
    m = Metrics.objects.filter(id=metrics_id).first() if metrics_id else None
    if not m:
        m = Metrics(user=_resolve_app_user(request))

    # Save Decimal fields safely
    decimal_fields = [
        "gifa_m2",
        "external_wall_area_m2",
        "external_openings_m2",
        "building_footprint_m2",
        "roof_area_m2",
        "roof_percent_gifa",
        "basement_size_m2",
        "basement_percent_gifa",
    ]
    for field in decimal_fields:
        if hasattr(m, field):
            setattr(m, field, _to_dec(payload.get(field), default=Decimal("0")))

    # Save Integer fields
    if hasattr(m, "num_apartments"):
        m.num_apartments = _to_int(payload.get("num_apartments"))
    if hasattr(m, "num_keys"):
        m.num_keys = _to_int(payload.get("num_keys"))
    if hasattr(m, "num_wcs"):
        m.num_wcs = _to_int(payload.get("num_wcs"))

    # Save Boolean and string fields
    if hasattr(m, "basement_present"):
        bp = payload.get("basement_present")
        m.basement_present = str(bp).lower() in ("1", "true", "yes", "on")
    if hasattr(m, "building_type"):
        m.building_type = payload.get("building_type") or getattr(m, "building_type", None)

    # Save total budget safely
    if hasattr(m, "total_budget_aud"):
        global_budget = payload.get("global_budget")
        m.total_budget_aud = _to_dec(global_budget, default=Decimal("0"))

    # Ensure ownership
    if not m.user:
        m.user = _resolve_app_user(request)

    # Save metrics and handle errors
    try:
        m.save()
    except Exception as e:
        logger.exception("Failed to save Metrics")
        return JsonResponse({"ok": False, "error": f"{e.__class__.__name__}: {e}"}, status=400)

    # Update session with current Metrics
    request.session["metrics_id"] = m.id
    request.session.modified = True

    return JsonResponse({"ok": True, "metrics_id": m.id})


# =========================
# Carbon / Calculator Views
# =========================

@login_required(login_url='login')
def carbon_view(request):
    """
    Render the carbon page with all interventions grouped by theme/class.
    Converts DB interventions to JSON for front-end use.
    """
    interventions = Interventions.objects.all()
    interventions_dict = {}

    # Reverse lookup for class aliases
    CLASS_ALIASES_REVERSE = {}
    for key, aliases in CLASS_ALIASES.items():
        for a in aliases:
            CLASS_ALIASES_REVERSE[a.lower()] = key

    # Group interventions by theme/class
    for i in interventions:
        db_theme = (i.theme or "other").lower()
        cls_key = CLASS_ALIASES_REVERSE.get(db_theme, "other")
        interventions_dict.setdefault(cls_key, []).append(
            {
                "id": i.id,
                "name": i.name or f"Intervention #{i.id}",
                "cost": float(i.cost_level or 0),
                "rating": float(i.intervention_rating or 0),
                "badges": [i.theme.capitalize()] if i.theme else [],
            }
        )

    # Define display classes for front-end
    classes = [
        {"key": "carbon", "label": "Carbon", "target": 80},
        {"key": "health", "label": "Health & Wellbeing", "target": 60},
        {"key": "water", "label": "Water Use", "target": 30},
        {"key": "circular", "label": "Circular Economy", "target": 40},
        {"key": "resilience", "label": "Resilience", "target": 60},
        {"key": "value", "label": "Value & Cost", "target": 10},
        {"key": "biodiversity", "label": "Biodiversity", "target": 20},
        {"key": "other", "label": "Other", "target": 0},
    ]

    return render(
        request,
        "carbon.html",
        {"interventions_json": json.dumps(interventions_dict), "classes": classes},
    )


@login_required(login_url='login')
def get_intervention_effects(request):
    """
    Given a source intervention name, return adjusted effects on target interventions.
    """
    source_name = request.GET.get("source")
    if not source_name:
        return JsonResponse({"error": "No source provided"}, status=400)

    effects = InterventionEffects.objects.filter(source_intervention_name=source_name)
    data = []

    for e in effects:
        target = Interventions.objects.filter(name=e.target_intervention_name).first()
        if not target:
            continue

        base_rating = float(target.intervention_rating or 0)
        max_effect_percent = 0.2  # ±20% max
        if e.effect_value is not None:
            effect_factor = float(e.effect_value) / 10 * max_effect_percent
            adjusted_rating = base_rating * (1 + effect_factor)
        else:
            adjusted_rating = base_rating

        data.append(
            {
                "target": e.target_intervention_name,
                "effect": round(adjusted_rating, 2),
                "note": e.note,
            }
        )

    return JsonResponse({"effects": data})


@login_required(login_url='login')
def calculator(request: HttpRequest):
    """
    Render calculator view or process calculator POST requests.
    GET -> render class targets
    POST -> delegate to _process_calculator_post
    """
    if request.method == "GET":
        class_targets = list(ClassTargets.objects.values("class_name", "target_rating"))
        return render(request, "calculator.html", {"class_targets": class_targets})
    return _process_calculator_post(request)


def intervention_effects(
    metric, interventions, selected_ids: Optional[List[int]] = None
):
    """
    Calculate adjusted intervention ratings based on dependencies, stage, and selection.
    Returns interventions grouped by class/theme.
    """
    grouped_interventions = {}
    max_stage = 0

    # Determine max stage among selected interventions
    if selected_ids:
        for i in interventions:
            if i.id in selected_ids:
                stage_val = getattr(i, "stage", 0) or 0
                max_stage = max(max_stage, int(stage_val))

    for i in interventions:
        # Skip interventions below max stage if selection exists
        stage_val = getattr(i, "stage", 0) or 0
        if selected_ids and stage_val < max_stage:
            continue

        # Dependency checks
        include = True
        deps = InterventionDependencies.objects.filter(intervention_id=i.id)
        for dep in deps:
            val = getattr(metric, dep.metric_name, None)
            if val is None:
                continue
            try:
                val = Decimal(val)
            except Exception:
                continue
            if (dep.min_value is not None and val < dep.min_value) or (
                dep.max_value is not None and val > dep.max_value
            ):
                include = False
                break
        if not include:
            continue

        # Base rating adjustment
        adjusted_rating = float(i.intervention_rating or 0)
        if selected_ids and i.id in selected_ids:
            adjusted_rating *= 1.1  # +10% rating for selected interventions

        cls = i.theme or "Other"
        grouped_interventions.setdefault(cls, []).append(
            {
                "id": str(i.id),
                "name": i.name or f"Intervention #{i.id}",
                "cost_level": float(i.cost_level or 0),
                "intervention_rating": round(adjusted_rating, 2),
                "description": i.description or "No description available",
                "stage": stage_val,
                "class_name": getattr(i, "class_name", ""),
                "theme": cls,
            }
        )

    return grouped_interventions



def _process_calculator_post(request: HttpRequest) -> HttpResponse:
    """
    Handle POST requests for the calculator page.
    Reads selected interventions from form data or JSON body,
    applies stage filtering, calculates adjusted ratings,
    and renders the results template.
    """
    # Get the current project metrics for this user/session
    metric = _get_current_metric(request)

    # Attempt to read selected intervention IDs from form POST
    try:
        selected_ids = request.POST.getlist("selected_ids") or []
    except Exception:
        selected_ids = []

    # If not found in form, try parsing JSON body
    if not selected_ids:
        try:
            payload = json.loads(request.body.decode("utf-8") or "{}")
            selected_ids = payload.get("selected_ids") or []
        except Exception:
            selected_ids = []

    # Ensure all IDs are integers
    selected_ids = [int(x) for x in selected_ids if str(x).strip().isdigit()]

    # Get all interventions from DB
    interventions_qs = list(Interventions.objects.all())

    # If user selected interventions, filter by max stage
    if selected_ids:
        max_stage = max(
            [
                getattr(i, "stage", 0) or 0
                for i in interventions_qs
                if i.id in selected_ids
            ]
        )
        # Only include interventions at or above the max stage
        interventions_qs = [
            i for i in interventions_qs if (getattr(i, "stage", 0) or 0) >= max_stage
        ]

    # Calculate grouped interventions with adjusted ratings
    grouped = intervention_effects(metric, interventions_qs, selected_ids)

    # Render the results page
    return render(
        request,
        "calculator_results.html",
        {
            "interventions": grouped,
            "interventions_json": json.dumps(grouped),
            "classes": [
                {"key": "carbon", "label": "Carbon", "target": 80},
                {"key": "health", "label": "Health & Wellbeing", "target": 60},
                {"key": "water", "label": "Water Use", "target": 30},
                {"key": "circular", "label": "Circular Economy", "target": 40},
                {"key": "resilience", "label": "Resilience", "target": 60},
                {"key": "value", "label": "Value & Cost", "target": 10},
                {"key": "biodiversity", "label": "Biodiversity", "target": 20},
            ],
            "cap_high": 300000,
            # Include current project ID for frontend API calls
            "metrics_id": metric.id,
        },
    )


# =========================
# Intervention Selection APIs
# =========================

@require_GET
@login_required(login_url='login')
def intervention_selection_list_api(request, metrics_id: int):
    """
    Return all interventions with a boolean 'selected' flag
    for the specified Metrics project.
    """
    # Get the Metrics project or 404
    project = get_object_or_404(Metrics, pk=metrics_id)

    # Get all currently selected intervention IDs for this project
    selected_ids = set(
        InterventionSelection.objects
        .filter(project=project)
        .values_list("intervention_id", flat=True)
    )

    # Build response list with selection status
    items = []
    for i in Interventions.objects.all().order_by("theme", "name"):
        items.append({
            "id": i.id,
            "name": i.name or f"Intervention #{i.id}",
            "theme": i.theme or "",
            "description": i.description or "",
            "cost_level": float(i.cost_level or 0),
            "intervention_rating": float(i.intervention_rating or 0),
            "selected": i.id in selected_ids,
        })

    return JsonResponse({"items": items, "project_id": project.id})


@require_POST
@login_required(login_url='login')
def intervention_selection_save_api(request, metrics_id: int):
    """
    Save the selected interventions for a Metrics project.
    The DB will exactly match the provided list of selected_ids.
    Expects a JSON body: {"selected_ids": [1,2,3,...]}
    """
    project = get_object_or_404(Metrics, pk=metrics_id)

    # Parse JSON payload
    try:
        payload = json.loads(request.body.decode("utf-8") or "{}")
    except json.JSONDecodeError:
        return HttpResponseBadRequest("Invalid JSON payload")

    selected_ids = payload.get("selected_ids")
    if not isinstance(selected_ids, list):
        return HttpResponseBadRequest("selected_ids must be a list")

    try:
        # Ensure all IDs are integers
        selected_ids = {int(x) for x in selected_ids if str(x).isdigit()}
    except Exception:
        return HttpResponseBadRequest("selected_ids must contain integers")

    # Existing selected interventions in DB
    existing_ids = set(
        InterventionSelection.objects
        .filter(project=project)
        .values_list("intervention_id", flat=True)
    )

    # Determine additions and deletions
    to_add = selected_ids - existing_ids
    to_del = existing_ids - selected_ids

    app_user = _resolve_app_user(request)

    # Apply DB changes atomically
    with transaction.atomic():
        if to_del:
            InterventionSelection.objects.filter(project=project, intervention_id__in=to_del).delete()
        if to_add:
            rows = [
                InterventionSelection(
                    project=project,
                    intervention_id=iid,
                    selected_by=app_user,  # Record the app user who selected
                )
                for iid in to_add
            ]
            InterventionSelection.objects.bulk_create(rows, ignore_conflicts=True)

    # Return updated selection info
    return JsonResponse({
        "ok": True,
        "added": sorted(to_add),
        "removed": sorted(to_del),
        "project_id": project.id,
        "total_selected": InterventionSelection.objects.filter(project=project).count(),
    })


# =========================
# Project List / Detail Views
# =========================

@login_required(login_url='login')
def projects_view(request: HttpRequest):
    """
    Display all projects in Metrics table.
    Supports optional search filtering by name, type, or location.
    """
    q = (request.GET.get("q") or "").strip()
    qs = Metrics.objects.all().order_by("-updated_at", "-created_at")

    if q:
        qs = qs.filter(
            Q(project_name__icontains=q)
            | Q(building_type__icontains=q)
            | Q(location__icontains=q)
        )

    return render(request, "projects.html", {"projects": qs, "query": q})


@login_required(login_url='login')
def project_detail_view(request, pk: int):
    """
    View or edit a project's details.
    GET with ?edit=1 enables editable mode.
    POST saves submitted changes.
    """
    p = get_object_or_404(Metrics, id=pk)

    if request.method == "POST":
        # Update string fields
        p.project_name = (request.POST.get("project_name") or p.project_name or "").strip()
        p.location = (request.POST.get("location") or p.location or "").strip()
        p.building_type = (request.POST.get("building_type") or p.building_type or "").strip()

        # Update decimal fields
        for f in [
            "gifa_m2",
            "external_wall_area_m2",
            "external_openings_m2",
            "building_footprint_m2",
            "roof_area_m2",
            "roof_percent_gifa",
            "basement_size_m2",
            "basement_percent_gifa",
        ]:
            if hasattr(p, f):
                setattr(p, f, _to_dec(request.POST.get(f), default=Decimal("0")))

        # Update integer and boolean fields
        p.num_apartments = _to_int(request.POST.get("num_apartments"))
        p.num_keys = _to_int(request.POST.get("num_keys"))
        p.num_wcs = _to_int(request.POST.get("num_wcs"))
        p.basement_present = bool(request.POST.get("basement_present"))

        p.save()

        # Store as current project in session
        request.session["metrics_id"] = p.id
        request.session.modified = True

        # Redirect to interventions page if requested
        if request.POST.get("next") == "interventions":
            return redirect("carbon")

        # Redirect back to detail page in edit mode with saved flag
        return redirect(f"{reverse('project_detail', args=[p.id])}?edit=1&saved=1")

    # GET request: display project details
    can_edit = request.GET.get("edit") == "1"
    request.session["metrics_id"] = p.id
    request.session.modified = True
    return render(request, "project_detail.html", {"p": p, "can_edit": can_edit})


# =========================
# Authentication + User Settings
# =========================

def login_view(request: HttpRequest):
    """
    Render login page and handle login submissions.
    """
    if request.user.is_authenticated:
        return redirect("home")
    if request.method == "POST":
        user = authenticate(
            request,
            username=request.POST.get("username"),
            password=request.POST.get("password"),
        )
        if user:
            login(request, user)
            return redirect("home")
        messages.error(request, "Invalid username or password.")
    return render(request, "login.html")


def logout_view(request: HttpRequest):
    """
    Logout user and redirect to login page.
    """
    if request.method == "POST":
        logout(request)
        messages.success(request, "Logged out successfully.")
        return redirect("login")
    return redirect("home")


def register_view(request: HttpRequest):
    """
    Handle user registration. Creates user and logs in upon success.
    """
    if request.method == "POST":
        username = request.POST.get("username")
        email = request.POST.get("email")
        password1 = request.POST.get("password1")
        password2 = request.POST.get("password2")

        if not all([username, email, password1, password2]):
            messages.error(request, "All fields are required.")
        elif password1 != password2:
            messages.error(request, "Passwords do not match.")
        elif User.objects.filter(username=username).exists():
            messages.error(request, "Username already exists.")
        elif User.objects.filter(email=email).exists():
            messages.error(request, "Email already exists.")
        else:
            user = User.objects.create_user(username=username, email=email, password=password1)
            login(request, user)
            messages.success(request, "Registration successful!")
            return redirect("home")
    return render(request, "register.html")


@login_required
def settings_view(request):
    """
    Display and handle user settings page.
    Includes theme selection, profile info update, and password change.
    """
    user = request.user
    current_theme = request.session.get('theme', 'light')

    if request.method == 'POST':
        try:
            # Theme change
            if 'theme_select' in request.POST:
                new_theme = request.POST.get('theme_select', 'light')
                request.session['theme'] = new_theme
                request.session.modified = True
                messages.success(request, f"Theme changed to {new_theme} mode!")
                return redirect('settings')

            # Profile info update
            if 'update_profile' in request.POST:
                new_username = request.POST.get('username')
                new_email = request.POST.get('email')

                if not new_username or not new_email:
                    raise ValueError("Username and email cannot be blank.")

                if User.objects.filter(username=new_username).exclude(id=user.id).exists():
                    raise ValueError("Username already exists.")
                if User.objects.filter(email=new_email).exclude(id=user.id).exists():
                    raise ValueError("Email already exists.")

                user.username = new_username
                user.email = new_email
                user.save()
                messages.success(request, "Profile updated successfully!")

            # Password change
            if 'change_password' in request.POST:
                current_password = request.POST.get('current_password')
                new_password = request.POST.get('new_password')
                confirm_password = request.POST.get('confirm_password')

                if not current_password or not new_password or not confirm_password:
                    raise ValueError("All password fields are required.")

                if new_password != confirm_password:
                    raise ValueError("New passwords do not match.")

                if not user.check_password(current_password):
                    raise ValueError("Current password is incorrect.")

                user.set_password(new_password)
                user.save()
                from django.contrib.auth import update_session_auth_hash
                update_session_auth_hash(request, user)
                messages.success(request, "Password changed successfully!")

        except ValueError as ve:
            messages.error(request, f"Error: {ve}")
        except Exception as e:
            messages.error(request, "Unexpected error occurred. Please try again later.")
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f"Settings update failed: {e}")

        return redirect('settings')

    # GET request: render settings page
    context = {'user': user, 'current_theme': current_theme}
    return render(request, 'settings.html', context)


@login_required(login_url='login')
def home(request):
    """
    Render home page/dashboard.
    """
    return render(request, "home.html")


@login_required(login_url='login')
def calculator_results(request):
    """
    Render calculator results page for a selected theme/class.
    """
    cls = request.GET.get("cls", "carbon")
    interventions_qs = Interventions.objects.filter(theme=cls).order_by(
        "-intervention_rating", "cost_level"
    )

    interventions = []
    for i in interventions_qs:
        interventions.append({
            "id": str(i.id),
            "name": i.name,
            "theme": i.theme,
            "description": i.description,
            "cost_level": float(i.cost_level or 0),
            "intervention_rating": float(i.intervention_rating or 0),
            "cost_range": getattr(i, "cost_range", ""),
        })

    metric = _get_current_metric(request)

    return render(
        request,
        "calculator_results.html",
        {
            "interventions_json": json.dumps(interventions),
            "classes": [
                {"key": "carbon", "label": "Carbon", "target": 80},
                {"key": "health", "label": "Health & Wellbeing", "target": 60},
                {"key": "water", "label": "Water Use", "target": 30},
                {"key": "circular", "label": "Circular Economy", "target": 40},
                {"key": "resilience", "label": "Resilience", "target": 60},
                {"key": "value", "label": "Value & Cost", "target": 10},
                {"key": "biodiversity", "label": "Biodiversity", "target": 20},
            ],
            "cap_high": 300000,
            "metrics_id": metric.id,
        },
    )



@login_required(login_url='login') 
def dashboard_view(request: HttpRequest):
    # --- Projects / budgets ---
    # Get the 3 most recently updated or created projects
    latest_projects = Metrics.objects.order_by("-updated_at", "-created_at")[:3]
    # Count total number of projects
    total_projects = Metrics.objects.count()

    # Calculate average budget across all projects
    avg_budget = (
        Metrics.objects.aggregate(avg_budget=Avg("total_budget_aud"))["avg_budget"]
        or Decimal("0")
    )

    # --- Intervention stats ---
    # Calculate average intervention rating across all interventions
    avg_intervention_rating = (
        Interventions.objects.aggregate(avg_rating=Avg("intervention_rating"))["avg_rating"]
        or 0
    )

    # Find the theme with the highest average intervention rating
    top_theme_data = (
        Interventions.objects.values("theme")
        .annotate(avg_rating=Avg("intervention_rating"))
        .order_by("-avg_rating")
        .first()
    )
    top_theme = top_theme_data["theme"] if top_theme_data else "N/A"
    top_theme_rating = round(top_theme_data["avg_rating"], 2) if top_theme_data else 0

    # --- Year-over-year projects statistics (last 6 years including current) ---
    now = timezone.now()
    start_year = now.year - 5

    # Count projects per year for the last 6 years
    yoy_raw = (
        Metrics.objects.annotate(y=ExtractYear("created_at"))
        .filter(y__gte=start_year, y__lte=now.year)
        .values("y")
        .annotate(n=Count("id"))
        .order_by("y")
    )

    # Map year to project count
    yoy_map = {row["y"]: int(row["n"]) for row in yoy_raw}
    # Create labels for last 6 years
    yoy_labels = [str(y) for y in range(start_year, now.year + 1)]
    # Get counts corresponding to labels, default to 0 if missing
    yoy_counts = [yoy_map.get(int(lbl), 0) for lbl in yoy_labels]

    # Prepare context for template
    context = {
        "latest_projects": latest_projects,
        "total_projects": total_projects,
        "avg_budget": avg_budget,
        "avg_intervention_rating": avg_intervention_rating,
        "top_theme": top_theme,
        "top_theme_rating": top_theme_rating,
        "yoy_labels_json": json.dumps(yoy_labels),
        "yoy_counts_json": json.dumps(yoy_counts),
        "current_year": now.year,
    }

    return render(request, "dashboard.html", context)


@login_required(login_url='login')
def carbon_2_view(request):
    # Render the Carbon 2 page
    return render(request, "carbon_2.html")


@login_required(login_url='login')
def reports_page(request: HttpRequest):
    """
    Reports overview page - shows all projects that can generate reports
    """
    # Resolve the current application user
    user = _resolve_app_user(request)
    if user:
        # If user exists, get projects associated with user
        projects = Metrics.objects.filter(user=user).order_by("-updated_at", "-created_at")
    else:
        # If no user, fallback to projects stored in session
        session_ids = request.session.get("my_project_ids", [])
        projects = Metrics.objects.filter(id__in=session_ids).order_by("-updated_at", "-created_at")
    
    return render(request, "reports.html", {
        "projects": projects
    })


@login_required(login_url='login')
def generate_report(request: HttpRequest, project_id: int):
    """
    Generate a report for a specific project based on download type
    """
    # Get the project or return 404 if not found
    project = get_object_or_404(Metrics, id=project_id)
    
    # Check if user has access to this project
    user = _resolve_app_user(request)
    if user and project.user != user:
        session_ids = request.session.get("my_project_ids", [])
        if project.id not in session_ids:
            return redirect("reports")
    
    # Determine report format
    download_format = request.GET.get('download')
    
    if download_format == 'pdf':
        return _generate_pdf_report(project)
    elif download_format == 'word':
        return _generate_word_report(project)
    else:
        return _generate_html_report(request, project)


def _generate_html_report(request: HttpRequest, project: Metrics):
    """
    Generate HTML report for a project including interventions and summary statistics
    """
    # Get all intervention selections for this project
    selections = InterventionSelection.objects.filter(project_id=project.id)
    intervention_ids = []

    # Extract intervention IDs from selections
    for s in selections:
        if hasattr(s, 'intervention_id') and s.intervention_id:
            intervention_ids.append(s.intervention_id)
    
    # Get the interventions to display in the report
    if intervention_ids:
        selected_interventions = list(Interventions.objects.filter(id__in=intervention_ids))
    else:
        # Fallback to first 5 interventions if none selected
        selected_interventions = list(Interventions.objects.all()[:5])
    
    # Calculate theme-level statistics for table display
    theme_stats = {}
    for intervention in selected_interventions:
        theme = intervention.theme or 'Other'
        rating = float(intervention.intervention_rating or 0)
        cost_level = float(intervention.cost_level or 0)
        
        if theme not in theme_stats:
            theme_stats[theme] = {
                'count': 0,
                'total_rating': 0,
                'total_cost': 0,
                'interventions': []
            }
        
        # Accumulate counts, ratings, costs, and list of interventions per theme
        theme_stats[theme]['count'] += 1
        theme_stats[theme]['total_rating'] += rating
        theme_stats[theme]['total_cost'] += cost_level
        theme_stats[theme]['interventions'].append(intervention)
    
    # Create aggregated table data for template
    available_interventions_data = []
    for theme, data in theme_stats.items():
        avg_rating = data['total_rating'] / data['count'] if data['count'] > 0 else 0
        avg_cost = data['total_cost'] / data['count'] if data['count'] > 0 else 0
        
        available_interventions_data.append({
            'theme': theme,
            'count': data['count'],
            'avg_rating': round(avg_rating, 1),
            'avg_cost': round(avg_cost, 1)
        })
    
    # Create context including all variable names used in templates
    context = {
        'project': project,
        
        # Provide multiple references for the same interventions list for template flexibility
        'selected_interventions': selected_interventions,
        'interventions': selected_interventions,
        'recommended_interventions': selected_interventions,
        'interventions_list': selected_interventions,
        'all_interventions': selected_interventions,
        
        # Table data for themes
        'available_interventions': available_interventions_data,
        'intervention_stats': available_interventions_data,
        
        # Theme-level statistics
        'theme_impacts': theme_stats,
        
        # Project summary information
        'metrics_summary': {
            'building_type': project.building_type or 'Not specified',
            'location': project.location or 'Not specified', 
            'total_area': f"{project.gifa_m2 or 0} m²",
            'total_budget': f"${project.total_budget_aud or 0:,.2f}",
            'apartments': project.num_apartments or 0,
            'basement': 'Yes' if project.basement_present else 'No',
            'created_date': project.created_at.strftime("%B %d, %Y"),
        },
        
        # Current report date and total selected interventions
        'report_date': timezone.now().strftime("%B %d, %Y"),
        'total_selected': len(selected_interventions),
    }
    
    # Render the HTML report template
    return render(request, "report_template.html", context)


def _generate_pdf_report(project: Metrics):
    """
    Generate PDF report (placeholder implementation)
    """
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="project_{project.id}_report.pdf"'
    
    # Placeholder PDF content
    response.write(b'PDF report generation would go here')
    return response


def _generate_word_report(project: Metrics):
    """
    Build a .docx report with the same data you show in the HTML report.
    """
    # ---- reuse your existing data builder (same as in _generate_html_report) ----
    # If you don’t have a helper, copy the aggregation from _generate_html_report here.
    # Below I inline a small version that matches your current context shape.

    # Selected interventions
    selections = InterventionSelection.objects.filter(project_id=project.id)
    ids = [s.intervention_id for s in selections if getattr(s, "intervention_id", None)]
    if ids:
        selected = list(Interventions.objects.filter(id__in=ids))
    else:
        selected = list(Interventions.objects.all()[:5])

    # Theme stats
    theme_stats = {}
    for iv in selected:
        theme = iv.theme or "Other"
        rating = float(iv.intervention_rating or 0)
        cost   = float(iv.cost_level or 0)
        bucket = theme_stats.setdefault(theme, {"count":0, "total_rating":0, "total_cost":0})
        bucket["count"] += 1
        bucket["total_rating"] += rating
        bucket["total_cost"] += cost

    # ---- build the document ----
    doc = Document()

    h = doc.add_heading('Environmental Impact Report', level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Project meta
    meta = doc.add_paragraph()
    meta.add_run(f'Project: {project.project_name}').bold = True
    meta.add_run(f'  •  Location: {project.location or "Not specified"}')
    doc.add_paragraph(f'Building Type: {project.building_type or "Not specified"}')
    doc.add_paragraph(f'Total Area: {project.gifa_m2 or 0} m²')
    doc.add_paragraph(f'Total Budget: ${project.total_budget_aud or 0:,.2f}')
    doc.add_paragraph()

    # Sustainability Action Plan (theme table)
    doc.add_heading('Sustainability Action Plan', level=1)
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = 'Focus Area'
    hdr[1].text = 'Actions'
    hdr[2].text = 'Avg Impact'
    hdr[3].text = 'Avg Investment'

    for theme, agg in theme_stats.items():
        count = max(1, agg["count"])
        avg_rating = round(agg["total_rating"] / count, 1)
        avg_cost   = round(agg["total_cost"] / count, 1)
        row = table.add_row().cells
        row[0].text = str(theme)
        row[1].text = f'{agg["count"]} measures'
        row[2].text = f'{avg_rating}/10'
        row[3].text = f'Level {avg_cost}'

    doc.add_paragraph()

    # Recommended Implementation Plan
    doc.add_heading('Recommended Implementation Plan', level=1)
    if selected:
        for iv in selected:
            doc.add_heading(iv.name, level=2)
            if getattr(iv, "description", None):
                doc.add_paragraph(str(iv.description))
            meta = doc.add_paragraph()
            meta.add_run('Effectiveness: ').bold = True
            meta.add_run(f'{getattr(iv, "intervention_rating", 0)}/10   ')
            meta.add_run('Implementation Cost: ').bold = True
            meta.add_run(f'Level {getattr(iv, "cost_level", 0)}/5')
            doc.add_paragraph()
    else:
        doc.add_paragraph('No specific recommendations selected.')

    # ---- return as download ----
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="project_{project.id}_report.docx"'
    return response
